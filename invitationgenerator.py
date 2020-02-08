import docx
def invitel(fname):
    doc.add_paragraph('It would be a pleasure to have the company of'.rjust(200)
    doc.add_paragraph(name)
    doc.add('at 11010 Memory Lane on the Evening of',)
    doc.add_paragraph('April 1st')
    doc.add_paragraph("at 7 o'clock".rjust(300))
    doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)

doc=docx.Document('inviteformat.docx')
guest=open('guest.txt','r')
guestt=guest.readlines()
for name in guestt:
    invite(name)
doc.save('inviteletter.docx')
